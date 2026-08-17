# -*- coding: utf-8 -*-
"""결과 문서 생성: DOCX(하이퍼링크 포함)/TXT/RIS/BibTeX, 폴더 일괄 종합 리포트."""
import io
import re
import time

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_ACCENT = RGBColor(0x1F, 0x4E, 0x79)
_WARN = RGBColor(0xB0, 0x3A, 0x2E)
_OK = RGBColor(0x2E, 0x7D, 0x32)
_AMBER = RGBColor(0x9A, 0x67, 0x00)

STATUS_LABEL = {
    "verified": "실존 확인", "mismatch": "제목 불일치", "not_found": "미발견",
    "suspect": "실존 의심", "link_ok": "링크 정상", "link_dead": "링크 오류",
    "skipped": "대조 생략",
}


def _setup(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _h(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.color.rgb = _ACCENT
    run.font.name = "맑은 고딕"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    return p


def _add_hyperlink(paragraph, url: str, text: str, size=None):
    """python-docx 문단에 외부 하이퍼링크 삽입."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "1F4E79"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    if size:
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2))); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


def _para_with_links(doc, text: str):
    """URL 부분을 클릭 가능한 하이퍼링크로 만든 문단."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.first_line_indent = Pt(-20)
    pos = 0
    for m in re.finditer(r"https?://[^\s]+", text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        _add_hyperlink(p, m.group(0).rstrip(".,"), m.group(0))
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def _verify_detail_full(v: dict) -> str:
    parts = [v.get("detail", "")]
    if v.get("journal"):
        parts.append(v["journal"].get("detail", ""))
    if v.get("preprint"):
        parts.append(v["preprint"].get("detail", ""))
    return " · ".join(x for x in parts if x)


def _certificate_page(doc: Document, result: dict):
    """검사 확인서 1쪽 — 투고 시 첨부할 수 있는 요약 표지.

    KCI 문헌 유사도 검사 결과서처럼 '검사를 거쳤다'는 사실을 한 장으로
    증빙하는 용도다. 상세 내역은 다음 쪽부터의 본 보고서가 담당한다.
    """
    s = result.get("summary", {})
    items = result.get("items", [])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(48)
    run = title.add_run("참고문헌 검사 확인서")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Reference Check Certificate — refcheck.kr").font.size = Pt(10)
    doc.add_paragraph()

    rows = [
        ("원고 파일", result.get("filename", "")),
        ("검사 일시", result.get("checked_at", "") or time.strftime("%Y-%m-%d %H:%M")),
        ("적용 기준", result.get("style_name", "")),
        ("처리 엔진", result.get("engine_label", "")),
        ("검사 도구", "refcheck.kr 참고문헌 검증·작성 서비스"
                     + (f" (버전 {result['app_version']})" if result.get("app_version") else "")),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        kr = tbl.rows[i].cells[0].paragraphs[0].add_run(k)
        kr.bold = True
        tbl.rows[i].cells[1].paragraphs[0].add_run(v)
    doc.add_paragraph()

    _h(doc, "검사 결과 요약", level=2)
    verified = s.get("verified", 0)
    warn = sum(1 for it in items
               if (it.get("verify") or {}).get("status") in ("mismatch", "not_found", "link_dead"))
    skipped = sum(1 for it in items
                  if not (it.get("verify") or {}).get("status")
                  or (it.get("verify") or {}).get("status") == "skipped")
    rows2 = [
        ("추출 참고문헌", f"{s.get('total', 0)}건"),
        ("형식 정리(변경)", f"{s.get('changed', 0)}건"),
        ("실존 확인(통과)", f"{verified}건"),
        ("미확인·불일치 (서지 오류 또는 DB 미수록 — 허위 아님)", f"{warn}건"),
        ("실존 의심 (다중 국제 DB 미발견)", f"{s.get('suspect', 0)}건"),
        ("철회(Retraction) 문헌 인용", f"{s.get('retracted', 0)}건"),
        ("서지 교정 제안", f"{s.get('suggestions', 0)}건"),
        ("검증 대조 생략", f"{skipped}건"),
    ]
    tbl2 = doc.add_table(rows=len(rows2), cols=2)
    tbl2.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows2):
        tbl2.rows[i].cells[0].paragraphs[0].add_run(k)
        vr = tbl2.rows[i].cells[1].paragraphs[0].add_run(v)
        vr.bold = True
    doc.add_paragraph()

    if result.get("verify_enabled"):
        p = doc.add_paragraph()
        p.add_run("검증 정보원 — 국내: KCI(한국학술지인용색인) · 국립중앙도서관 서지정보 · "
                  "국회도서관 국가학술정보 / 해외: Crossref · OpenAlex · Semantic Scholar · "
                  "DataCite · DOAJ · URL 접속 확인").font.size = Pt(8.5)
    note = doc.add_paragraph()
    nr = note.add_run(
        "이 확인서는 refcheck.kr(참고문헌 검증·작성 서비스)가 위 원고의 참고문헌을 "
        "자동 검사한 결과의 요약입니다. '미확인'은 대조한 데이터베이스에서 찾지 못했다는 "
        "뜻이며 문헌이 존재하지 않는다는 판정이 아닙니다. 항목별 상세 내역은 다음 쪽의 "
        "본 보고서를 참조하십시오.")
    nr.font.size = Pt(8.5)
    nr.font.color.rgb = RGBColor(0x7C, 0x8B, 0x9B)
    doc.add_page_break()


def build_result_docx(result: dict) -> bytes:
    doc = Document()
    _setup(doc)

    # 1쪽: 검사 확인서(투고 첨부용 요약) — 상세는 2쪽부터
    _certificate_page(doc, result)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("참고문헌 검증·작성 결과")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _ACCENT

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"원고: {result.get('filename', '')}   |   기준: {result.get('style_name', '')}   |   "
        f"엔진: {result.get('engine_label', '')}   |   {time.strftime('%Y-%m-%d %H:%M')}"
    ).font.size = Pt(9)

    sec = 0
    items = result.get("items", [])
    s = result.get("summary", {})

    # ── 처리 요약
    sec += 1
    _h(doc, f"{sec}. 처리 요약")
    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = "Light Grid Accent 1"
    heads = ["추출 문헌", "형식 수정", "확인 필요", "실존 확인", "교정 제안", "본문 대조 이슈"]
    vals = [str(s.get("total", 0)), str(s.get("changed", 0)), str(s.get("needs_check", 0)),
            str(s.get("verified", 0)), str(s.get("suggestions", 0)), str(s.get("crosscheck_issues", 0))]
    for i, h in enumerate(heads):
        c = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
        c.bold = True
    for i, v in enumerate(vals):
        tbl.rows[1].cells[i].paragraphs[0].add_run(v)
    if s.get("retracted"):
        p = doc.add_paragraph()
        r = p.add_run(f"⚠ 철회(Retraction)된 문헌 인용 {s['retracted']}건 — 인용 유지 여부를 반드시 재검토하세요.")
        r.bold = True
        r.font.color.rgb = _WARN
    if s.get("suspect"):
        p = doc.add_paragraph()
        r = p.add_run(f"⚠ 다중 DB 미발견(실존 의심) 문헌 {s['suspect']}건 — AI 생성 인용·서지 오류 가능성을 확인하세요.")
        r.font.color.rgb = _WARN

    # ── 최종 목록
    sec += 1
    _h(doc, f"{sec}. 정리된 참고문헌 목록")
    doc.add_paragraph("아래 목록은 선택한 기준의 형식과 배열 순서로 정리된 최종 참고문헌입니다(링크 클릭 가능).").runs[0].font.size = Pt(9)
    cur_group = None
    for item in items:
        grp = item.get("group", "")
        if grp and grp != cur_group:
            cur_group = grp
            gp = doc.add_paragraph()
            gr = gp.add_run(f"[{grp}]")
            gr.bold = True
        _para_with_links(doc, item.get("formatted", ""))
        v = item.get("verify") or {}
        badges = []
        if (v.get("retraction") or {}).get("severe"):
            badges.append(f"⚠ {v['retraction']['label']} 문헌")
        if v.get("status") == "suspect":
            badges.append("⚠ 실존 의심(다중 DB 미발견)")
        if item.get("issues"):
            badges.append(" / ".join(item["issues"]))
        if badges:
            ip = doc.add_paragraph()
            ip.paragraph_format.left_indent = Pt(20)
            ir = ip.add_run("· " + " · ".join(badges))
            ir.font.size = Pt(8.5)
            ir.font.color.rgb = _WARN
        for tp in item.get("tips") or []:
            tp_p = doc.add_paragraph()
            tp_p.paragraph_format.left_indent = Pt(20)
            tr = tp_p.add_run(f"💡 [{tp.get('label', '')}] {tp.get('rule', '')}"
                              + (f" (예: {tp['example']})" if tp.get("example") else ""))
            tr.font.size = Pt(8.5)
            tr.font.color.rgb = _AMBER

    # ── 대비표
    sec += 1
    _h(doc, f"{sec}. 변경 전·후 대비표")
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["#", "변경 전(원문)", "변경 후(기준 적용)"]):
        r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
    for i, item in enumerate(items, 1):
        row = tbl.add_row()
        row.cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(item.get("raw", "")).font.size = Pt(9)
        after = row.cells[2].paragraphs[0].add_run(item.get("formatted", ""))
        after.font.size = Pt(9)
        if item.get("changed"):
            after.font.color.rgb = _ACCENT

    # ── 실존·윤리 검증
    if result.get("verify_enabled"):
        sec += 1
        _h(doc, f"{sec}. 실존·윤리 검증 결과 (Crossref·OpenAlex·국내DB·철회검사)")
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "판정", "출처", "상세"]):
            r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
            r.bold = True
        for i, item in enumerate(items, 1):
            v = item.get("verify") or {}
            row = tbl.add_row()
            row.cells[0].paragraphs[0].add_run(str(i)).font.size = Pt(9)
            lab = row.cells[1].paragraphs[0].add_run(STATUS_LABEL.get(v.get("status", ""), "-"))
            lab.font.size = Pt(9)
            if v.get("status") in ("verified", "link_ok"):
                lab.font.color.rgb = _OK
            elif v.get("status") in ("mismatch", "not_found", "link_dead", "suspect"):
                lab.font.color.rgb = _WARN
            if (v.get("retraction") or {}).get("severe"):
                lab.font.color.rgb = _WARN
            row.cells[2].paragraphs[0].add_run(v.get("source", "") or "-").font.size = Pt(9)
            row.cells[3].paragraphs[0].add_run(_verify_detail_full(v)).font.size = Pt(9)

    # ── 교정 제안
    all_sugg = [(i, sg) for i, item in enumerate(items, 1)
                for sg in (item.get("suggestions") or [])]
    if all_sugg:
        sec += 1
        applied_any = any(sg.get("applied") for _, sg in all_sugg)
        _h(doc, f"{sec}. 서지요소 교정 " + ("적용 내역" if applied_any else "제안"))
        doc.add_paragraph(
            "검증 과정에서 매칭된 공식 서지(Crossref·KCI 등)와 원고 표기가 다른 항목입니다."
            + ("" if applied_any else " (자동 교정 옵션을 켜면 아래 제안이 자동 반영됩니다)")
        ).runs[0].font.size = Pt(9)
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "항목", "원고 표기", "공식 서지", "출처"]):
            r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
            r.bold = True
        for i, sg in all_sugg:
            row = tbl.add_row()
            for c, val in enumerate([str(i), sg.get("label", ""), sg.get("current", ""),
                                     sg.get("suggested", ""), sg.get("source", "")]):
                row.cells[c].paragraphs[0].add_run(val).font.size = Pt(9)

    # ── 작성 제안 (논문 사례를 통한 제안 · 발행본 검토 제안)
    all_tips = [(i, tp) for i, item in enumerate(items, 1)
                for tp in (item.get("tips") or [])]
    if all_tips:
        sec += 1
        _h(doc, f"{sec}. 참고문헌 작성 제안")
        doc.add_paragraph(
            "문편협 공통 기준이 일순위이며, 아래는 기준이 명시하지 않거나 애매한 부분에 대한 "
            "보조 제안입니다. '논문 사례를 통한 제안'은 4개 학회지의 2025년 이후 발행 논문 관행에서, "
            "'발행본 검토 제안'은 학회지 발행본과의 대조 검토에서 추출한 것입니다."
        ).runs[0].font.size = Pt(9)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(["#", "구분", "제안", "예시"]):
            r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
            r.bold = True
        for i, tp in all_tips:
            row = tbl.add_row()
            for c, val in enumerate([str(i), tp.get("label", ""), tp.get("rule", ""),
                                     tp.get("example", "")]):
                row.cells[c].paragraphs[0].add_run(val).font.size = Pt(9)

    # ── 본문 대조
    cc = result.get("crosscheck")
    if cc:
        sec += 1
        _h(doc, f"{sec}. 본문 인용 ↔ 참고문헌 목록 대조")
        doc.add_paragraph(f"본문에서 발견된 인용: {cc.get('citations_found', 0)}건").runs[0].font.size = Pt(9)
        if cc.get("cited_not_listed"):
            p = doc.add_paragraph()
            p.add_run("본문에 인용되었으나 참고문헌 목록에 없는 문헌").bold = True
            for c in cc["cited_not_listed"]:
                extra = " (같은 저자의 다른 연도 문헌은 목록에 있음 — 연도 확인)" if c.get("name_only_match") else ""
                doc.add_paragraph(f"· {c['name']}({c['year']}) — “…{c['snippet']}…”{extra}").runs[0].font.size = Pt(9)
        if cc.get("listed_not_cited"):
            p = doc.add_paragraph()
            p.add_run("참고문헌 목록에 있으나 본문에서 인용을 찾지 못한 문헌").bold = True
            for c in cc["listed_not_cited"]:
                doc.add_paragraph(f"· {c.get('authors') or ''} ({c.get('year', '')}) — {c.get('raw', '')}").runs[0].font.size = Pt(9)
        if not cc.get("cited_not_listed") and not cc.get("listed_not_cited"):
            doc.add_paragraph("이상 없음 — 본문 인용과 참고문헌 목록이 서로 일치합니다.")

    # ── 건전성 리포트
    hl = result.get("health")
    if hl:
        sec += 1
        _h(doc, f"{sec}. 참고문헌 건전성 리포트")
        n = hl.get("recent_years", 10)
        lines = [f"총 문헌 수: {hl.get('total', 0)}건"]
        if hl.get("year_min"):
            lines.append(f"발행연도 범위: {hl['year_min']} ~ {hl['year_max']}")
        if hl.get("recent_ratio") is not None:
            lines.append(f"최근 {n}년 문헌 비율: {hl['recent_ratio']:.0%} ({hl['recent_count']}건)")
        if hl.get("self_cites"):
            lines.append(f"자기인용(사용자 이름 일치): {hl['self_cites']}건")
        for line in lines:
            doc.add_paragraph("· " + line).runs[0].font.size = Pt(9.5)
        if hl.get("top_journals"):
            p = doc.add_paragraph()
            p.add_run("자주 인용된 학술지").bold = True
            for j in hl["top_journals"]:
                doc.add_paragraph(f"· {j['name']} — {j['count']}건 ({j['share']:.0%})").runs[0].font.size = Pt(9)
        if hl.get("duplicates"):
            p = doc.add_paragraph()
            r = p.add_run(f"중복 의심 참고문헌 {len(hl['duplicates'])}쌍")
            r.bold = True
            r.font.color.rgb = _WARN
            for d in hl["duplicates"]:
                doc.add_paragraph(f"· [{d['reason']}] {d['a']} ↔ {d['b']}").runs[0].font.size = Pt(9)

    # ── 영문 변환 목록
    eng = result.get("english_list")
    if eng:
        sec += 1
        _h(doc, f"{sec}. 국문 참고문헌의 영문 변환 목록 (알파벳순)")
        for line in eng:
            _para_with_links(doc, line)

    tail = doc.add_paragraph()
    tail.add_run(
        "\n※ 본 결과는 자동 변환·검증 결과이므로 투고 전 변경 전·후 대비표를 통해 최종 확인이 필요합니다. "
        "‘확인 필요’·‘실존 의심’·‘철회’ 표시 항목은 원문 데이터베이스(KCI, RISS, Crossref 등)에서 확인하세요."
    ).font.size = Pt(8.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_result_txt(result: dict) -> str:
    lines = [f"# 참고문헌 표준화 결과 — {result.get('filename', '')}",
             f"# 기준: {result.get('style_name', '')} | 엔진: {result.get('engine_label', '')}", ""]
    cur_group = None
    for item in result.get("items", []):
        grp = item.get("group", "")
        if grp and grp != cur_group:
            cur_group = grp
            lines.append(f"[{grp}]")
        lines.append(item.get("formatted", ""))
    tips = [(i, tp) for i, item in enumerate(result.get("items", []), 1)
            for tp in (item.get("tips") or [])]
    if tips:
        lines += ["", "# 작성 제안 (문편협 공통 기준 외 참고)"]
        for i, tp in tips:
            lines.append(f"[{i}] [{tp.get('label', '')}] {tp.get('rule', '')}"
                         + (f" (예: {tp['example']})" if tp.get("example") else ""))
    eng = result.get("english_list")
    if eng:
        lines += ["", "# 영문 변환 목록"]
        lines += eng
    return "\n".join(lines)


# ================================================================ RIS / BibTeX

_RIS_TYPE = {
    "journal": "JOUR", "book": "BOOK", "book_chapter": "CHAP", "thesis": "THES",
    "report": "RPRT", "newspaper": "NEWS", "web": "ELEC", "conference": "CPAPER",
    "law": "STAT", "standard": "STAND", "av": "ADVS", "interview": "PCOMM",
    "unknown": "GEN",
}


def build_ris(result: dict) -> str:
    out = []
    for item in result.get("items", []):
        e = item.get("entry") or {}
        t = item.get("type", "unknown")
        rec = [f"TY  - {_RIS_TYPE.get(t, 'GEN')}"]
        for a in e.get("authors") or []:
            if a:
                rec.append(f"AU  - {a}")
        year = re.sub(r"[a-z]$", "", e.get("year", "") or "")
        if re.fullmatch(r"\d{4}", year):
            rec.append(f"PY  - {year}")
        if e.get("title"):
            rec.append(f"TI  - {e['title']}")
        if e.get("container"):
            rec.append(f"T2  - {e['container']}")
        if e.get("volume"):
            rec.append(f"VL  - {e['volume']}")
        if e.get("issue"):
            rec.append(f"IS  - {e['issue']}")
        pages = e.get("pages", "")
        if pages and "-" in pages:
            sp, ep = pages.split("-", 1)
            rec.append(f"SP  - {sp.strip()}")
            rec.append(f"EP  - {ep.strip()}")
        elif pages:
            rec.append(f"SP  - {pages}")
        if t == "thesis":
            if e.get("institution"):
                rec.append(f"PB  - {e['institution']}")
            if e.get("degree"):
                rec.append(f"M3  - {e['degree']}")
        else:
            if e.get("publisher"):
                rec.append(f"PB  - {e['publisher']}")
            if e.get("place"):
                rec.append(f"CY  - {e['place']}")
        if e.get("doi"):
            rec.append(f"DO  - {e['doi']}")
        if e.get("url"):
            rec.append(f"UR  - {e['url']}")
        rec.append("ER  - ")
        out.append("\n".join(rec))
    return "\n\n".join(out) + "\n"


_BIB_TYPE = {
    "journal": "article", "book": "book", "book_chapter": "incollection",
    "thesis": "phdthesis", "report": "techreport", "conference": "inproceedings",
}


def _bib_escape(s: str) -> str:
    s = (s or "").replace("\\", "\\textbackslash{}")
    for ch, rep in (("{", "\\{"), ("}", "\\}"), ("&", "\\&"), ("%", "\\%"),
                    ("$", "\\$"), ("#", "\\#"), ("_", "\\_"),
                    ("~", "\\textasciitilde{}"), ("^", "\\textasciicircum{}")):
        s = s.replace(ch, rep)
    return re.sub(r"[\r\n]+", " ", s)


def _bib_clean_url(s: str) -> str:
    """doi/url 필드용: 중괄호·개행만 제거(LaTeX 이스케이프는 URL을 깨뜨림)."""
    return re.sub(r"[{}\r\n\s]+", "", s or "")


def build_bibtex(result: dict) -> str:
    out = []
    for i, item in enumerate(result.get("items", []), 1):
        e = item.get("entry") or {}
        t = item.get("type", "unknown")
        btype = _BIB_TYPE.get(t, "misc")
        if t == "thesis" and "석사" in (e.get("degree") or "") or "Master" in (e.get("degree") or ""):
            btype = "mastersthesis"
        fields = []
        authors = [a for a in (e.get("authors") or []) if a]
        if authors:
            fields.append(("author", " and ".join(_bib_escape(a) for a in authors)))
        if e.get("title"):
            fields.append(("title", _bib_escape(e["title"])))
        year = re.sub(r"[a-z]$", "", e.get("year", "") or "")
        if re.fullmatch(r"\d{4}", year):
            fields.append(("year", year))
        if e.get("container"):
            key = "journal" if btype == "article" else "booktitle" if btype in ("incollection", "inproceedings") else "howpublished"
            fields.append((key, _bib_escape(e["container"])))
        if e.get("volume"):
            fields.append(("volume", _bib_escape(e["volume"])))
        if e.get("issue"):
            fields.append(("number", _bib_escape(e["issue"])))
        if e.get("pages"):
            fields.append(("pages", _bib_escape(e["pages"].replace("-", "--"))))
        if t == "thesis":
            if e.get("institution"):
                fields.append(("school", _bib_escape(e["institution"])))
        else:
            if e.get("publisher"):
                fields.append(("publisher", _bib_escape(e["publisher"])))
            if e.get("place"):
                fields.append(("address", _bib_escape(e["place"])))
        if e.get("doi"):
            fields.append(("doi", _bib_clean_url(e["doi"])))
        if e.get("url"):
            fields.append(("url", _bib_clean_url(e["url"])))
        body = ",\n".join(f"  {k} = {{{v}}}" for k, v in fields)
        out.append(f"@{btype}{{ref{i},\n{body}\n}}")
    return "\n\n".join(out) + "\n"


# ================================================================ 종합 리포트

def build_batch_report_docx(results: list[dict], folder: str) -> bytes:
    doc = Document()
    _setup(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("참고문헌 표준화 종합 리포트")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _ACCENT
    doc.add_paragraph(f"대상 폴더: {folder}").runs[0].font.size = Pt(9)
    doc.add_paragraph(f"처리 일시: {time.strftime('%Y-%m-%d %H:%M')} | 처리 파일: {len(results)}건").runs[0].font.size = Pt(9)

    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(["파일", "문헌", "수정", "확인필요", "실존확인", "철회", "실존의심", "본문대조"]):
        r = tbl.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True
    for res in results:
        s = res.get("summary", {})
        row = tbl.add_row()
        vals = [res.get("filename", ""), s.get("total", 0), s.get("changed", 0),
                s.get("needs_check", 0), s.get("verified", 0), s.get("retracted", 0),
                s.get("suspect", 0), s.get("crosscheck_issues", 0)]
        for i, v in enumerate(vals):
            run = row.cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
            if i == 5 and v:
                run.font.color.rgb = _WARN
                run.bold = True

    for res in results:
        if res.get("error"):
            p = doc.add_paragraph()
            r = p.add_run(f"⚠ {res.get('filename', '')}: {res['error']}")
            r.font.color.rgb = _WARN
            r.font.size = Pt(9)

    doc.add_paragraph("\n파일별 상세 결과는 같은 폴더의 ‘(파일명)_참고문헌정리.docx’를 확인하세요.").runs[0].font.size = Pt(9)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
